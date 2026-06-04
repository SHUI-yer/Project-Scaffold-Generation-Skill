"""
docx2md.py — 将 .docx 文件转换为 Markdown 或纯文本
用法：python scripts/docx2md.py <input文件夹> [输出格式: md|txt]
依赖：pip install python-docx
"""
import sys
import os
from pathlib import Path

def convert_docx_to_markdown(docx_path: str) -> str:
    """将 .docx 转为 Markdown 格式"""
    try:
        from docx import Document
    except ImportError:
        print("错误：需要安装 python-docx")
        print("请运行：pip install python-docx")
        sys.exit(1)

    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""

        if style.startswith("Heading"):
            level = style.replace("Heading ", "")
            try:
                level = int(level)
            except ValueError:
                level = 1
            lines.append(f"{'#' * level} {para.text}")
        elif style.startswith("List"):
            lines.append(f"- {para.text}")
        else:
            text = para.text.strip()
            if text:
                lines.append(text)

    # 处理表格
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        if rows:
            # 表头
            header = rows[0]
            lines.append("")
            lines.append("| " + " | ".join(header) + " |")
            lines.append("|" + "|".join(["---"] * len(header)) + "|")
            # 数据行
            for row in rows[1:]:
                # 补齐列数
                while len(row) < len(header):
                    row.append("")
                lines.append("| " + " | ".join(row[:len(header)]) + " |")
            lines.append("")

    return "\n".join(lines)


def convert_docx_to_text(docx_path: str) -> str:
    """将 .docx 转为纯文本"""
    try:
        from docx import Document
    except ImportError:
        print("错误：需要安装 python-docx")
        print("请运行：pip install python-docx")
        sys.exit(1)

    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))

    return "\n".join(lines)


def main():
    if len(sys.argv) < 2:
        print("用法：python docx2md.py <input文件夹> [md|txt]")
        print("示例：python docx2md.py input/ md")
        sys.exit(1)

    input_dir = Path(sys.argv[1])
    output_format = sys.argv[2] if len(sys.argv) > 2 else "md"

    if not input_dir.exists():
        print(f"错误：目录不存在 {input_dir}")
        sys.exit(1)

    # 查找所有 .docx 文件
    docx_files = list(input_dir.glob("*.docx"))

    if not docx_files:
        print(f"在 {input_dir} 中未找到 .docx 文件")
        print("提示：.doc 文件需要先用 Word/WPS 另存为 .docx 格式")
        sys.exit(0)

    for docx_file in docx_files:
        print(f"转换：{docx_file.name}")

        if output_format == "md":
            content = convert_docx_to_markdown(str(docx_file))
            output_file = docx_file.with_suffix(".md")
        else:
            content = convert_docx_to_text(str(docx_file))
            output_file = docx_file.with_suffix(".txt")

        output_file.write_text(content, encoding="utf-8")
        print(f"  -> {output_file.name} ({len(content)} 字符)")

    print(f"\n完成！共转换 {len(docx_files)} 个文件")


if __name__ == "__main__":
    main()
